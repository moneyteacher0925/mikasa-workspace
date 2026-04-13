from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(11)

# Title
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('RS(Revenue Share) 제휴 파트너 계약서')
run.bold = True
run.font.size = Pt(18)

doc.add_paragraph()

# Intro
doc.add_paragraph(
    '주식회사 플로우머스(이하 "갑")와 _______________(이하 "을")는 '
    '상호 신뢰를 바탕으로 아래와 같이 수익 배분(Revenue Share) 제휴 파트너 계약을 체결한다.'
)

doc.add_paragraph()

# 제1조
h = doc.add_paragraph()
run = h.add_run('제1조 (목적)')
run.bold = True
p = doc.add_paragraph(
    '본 계약은 갑이 운영하는 오프라인 모임 플랫폼 "딱모여"를 통해 모집된 참가자를 대상으로, '
    '을이 자사 서비스/상품을 제공하고 발생하는 매출에 대해 수익을 배분하는 제휴 구조를 정하는 것을 목적으로 한다.'
)

doc.add_paragraph()

# 제2조
h = doc.add_paragraph()
run = h.add_run('제2조 (역할 분담)')
run.bold = True

doc.add_paragraph('1. 갑의 역할', style='List Number')
doc.add_paragraph('카테고리별 오프라인 모임 기획 및 주최', style='List Bullet')
doc.add_paragraph('모임 참가자 모집 (광고, 마케팅, 바이럴 등)', style='List Bullet')
doc.add_paragraph('모임 운영 및 현장 관리', style='List Bullet')

doc.add_paragraph('2. 을의 역할', style='List Number')
doc.add_paragraph('모임 현장 참석 및 서비스/상품 브리핑', style='List Bullet')
doc.add_paragraph('참가자 대상 영업 및 계약/결제 유도', style='List Bullet')
doc.add_paragraph('계약/결제 진행 현황에 대한 투명한 보고', style='List Bullet')

doc.add_paragraph()

# 제3조
h = doc.add_paragraph()
run = h.add_run('제3조 (수익 배분)')
run.bold = True

doc.add_paragraph(
    '1. 을의 서비스/상품 판매로 발생하는 매출에 대해 아래 비율로 수익을 배분한다.'
)

table = doc.add_table(rows=3, cols=2)
table.style = 'Table Grid'
table.alignment = WD_TABLE_ALIGNMENT.CENTER
cells = table.rows[0].cells
cells[0].text = '구분'
cells[1].text = '배분 비율'
for cell in cells:
    for p in cell.paragraphs:
        for r in p.runs:
            r.bold = True
cells = table.rows[1].cells
cells[0].text = '갑 (플로우머스)'
cells[1].text = '____%'
cells = table.rows[2].cells
cells[0].text = '을 (제휴 파트너)'
cells[1].text = '____%'

doc.add_paragraph()
doc.add_paragraph(
    '2. 수익 배분 기준 매출은 부가세를 포함한 총매출을 기준으로 한다. '
    '부가가치세 신고 및 납부 의무는 을에게 있다.'
)
doc.add_paragraph(
    '3. 환불, 취소, 환수 등으로 인한 매출 감소분은 을의 부담으로 한다. '
    '정산 완료 후 환수가 발생한 경우, 을은 차기 정산 시 해당 금액을 상계하거나 갑에게 반환한다.'
)
doc.add_paragraph(
    '4. 수익 배분금은 매월 말일 기준으로 정산하며, 익월 ____일까지 지급한다.'
)

doc.add_paragraph()

# 제4조
h = doc.add_paragraph()
run = h.add_run('제4조 (매출 추적 및 투명성)')
run.bold = True

doc.add_paragraph(
    '1. 온라인 결제의 경우: 갑이 제공하는 추적 링크를 통해 매출을 투명하게 추적하며, '
    '갑과 을 모두 실시간으로 데이터에 접근할 수 있도록 한다.'
)
doc.add_paragraph(
    '2. 오프라인 계약(대면 영업)의 경우: 을은 모임 참가자와의 미팅 후 진행 상황을 '
    '갑에게 보고하여야 하며, 갑이 제공하는 공유 대시보드(UI)를 통해 '
    '계약 진행 현황, 결제 여부, 금액 등을 투명하게 공유한다.'
)
doc.add_paragraph(
    '3. 을이 매출 정보를 성실하게 공유하지 않을 경우, 갑은 RS 구조를 중단하고 '
    '광고 제휴 방식으로 전환할 수 있다.'
)

doc.add_paragraph()

# 제5조
h = doc.add_paragraph()
run = h.add_run('제5조 (계약 기간)')
run.bold = True

doc.add_paragraph(
    '1. 본 계약의 유효기간은 체결일로부터 ____개월로 한다.'
)
doc.add_paragraph(
    '2. 계약 만료 ____일 전까지 쌍방 서면 이의가 없는 경우 동일 조건으로 자동 연장된다.'
)

doc.add_paragraph()

# 제6조
h = doc.add_paragraph()
run = h.add_run('제6조 (비밀유지)')
run.bold = True

doc.add_paragraph(
    '갑과 을은 본 계약의 내용 및 계약 이행 과정에서 취득한 상대방의 영업 비밀, '
    '고객 정보 등을 제3자에게 누설하거나 본 계약 목적 외에 사용하지 아니한다. '
    '본 조항은 계약 종료 후에도 유효하다.'
)

doc.add_paragraph()

# 제7조
h = doc.add_paragraph()
run = h.add_run('제7조 (계약 해지)')
run.bold = True

doc.add_paragraph('다음 각 호에 해당하는 경우 상대방에게 서면 통보 후 본 계약을 해지할 수 있다.')
doc.add_paragraph('상대방이 본 계약의 의무를 중대하게 위반한 경우', style='List Bullet')
doc.add_paragraph('매출 정보의 허위 보고 또는 은닉이 확인된 경우', style='List Bullet')
doc.add_paragraph('상대방의 파산, 회생 절차 개시 등 정상적인 영업이 불가능한 경우', style='List Bullet')

doc.add_paragraph()

# 제8조
h = doc.add_paragraph()
run = h.add_run('제8조 (분쟁 해결)')
run.bold = True

doc.add_paragraph(
    '본 계약과 관련하여 분쟁이 발생한 경우 갑의 본점 소재지를 관할하는 법원을 '
    '제1심 관할 법원으로 한다.'
)

doc.add_paragraph()
doc.add_paragraph()

# 서명란
doc.add_paragraph(
    '본 계약의 성립을 증명하기 위하여 계약서 2부를 작성하고, 갑과 을이 각각 서명 날인한 후 1부씩 보관한다.'
)

doc.add_paragraph()

date_p = doc.add_paragraph()
date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = date_p.add_run('2026년    월    일')
run.font.size = Pt(12)

doc.add_paragraph()
doc.add_paragraph()

# 갑
doc.add_paragraph('【갑】')
doc.add_paragraph('회사명: 주식회사 플로우머스')
doc.add_paragraph('대표자: ')
doc.add_paragraph('주  소: 서울특별시 서초구 서초대로 398 그레이츠강남 709호')
doc.add_paragraph('서  명: ______________________')

doc.add_paragraph()

# 을
doc.add_paragraph('【을】')
doc.add_paragraph('회사명: ')
doc.add_paragraph('대표자: ')
doc.add_paragraph('주  소: ')
doc.add_paragraph('서  명: ______________________')

out = '/Users/kwoneren/.openclaw/workspace/RS제휴파트너계약서_초안.docx'
doc.save(out)
print(f'Saved: {out}')
