from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

# 페이지 여백
for section in doc.sections:
    section.top_margin = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

style = doc.styles['Normal']
style.font.size = Pt(11)

def add_colored_heading(text, level=1, color=RGBColor(30, 58, 138)):
    h = doc.add_heading('', level=level)
    r = h.add_run(text)
    r.font.color.rgb = color
    return h

def para(text, bold=False, italic=False, color=None, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    if bold: r.bold = True
    if italic: r.italic = True
    if color: r.font.color.rgb = color
    if size: r.font.size = Pt(size)
    return p

def bullet(text, bold_prefix=None):
    p = doc.add_paragraph(style='List Bullet')
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        p.add_run(text)
    else:
        p.add_run(text)
    return p

def divider():
    p = doc.add_paragraph()
    r = p.add_run('─' * 50)
    r.font.color.rgb = RGBColor(200, 200, 200)
    r.font.size = Pt(8)

def tag_line(tag, content, tag_color=RGBColor(30, 58, 138)):
    p = doc.add_paragraph()
    r1 = p.add_run(f'  {tag}  ')
    r1.bold = True
    r1.font.color.rgb = RGBColor(255, 255, 255)
    # Can't do background in python-docx easily, use brackets instead
    p2 = doc.add_paragraph()
    r1 = p2.add_run(f'[{tag}] ')
    r1.bold = True
    r1.font.color.rgb = tag_color
    r1.font.size = Pt(11)
    r2 = p2.add_run(content)
    r2.font.size = Pt(11)
    # remove the first empty paragraph
    p._element.getparent().remove(p._element)
    return p2

# ═══════ 표지 ═══════
doc.add_paragraph()
doc.add_paragraph()
t = doc.add_heading('', level=0)
t.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = t.add_run('엠타트업 22기 — 4조')
r.font.size = Pt(28)
r.font.color.rgb = RGBColor(30, 58, 138)

t2 = doc.add_heading('', level=1)
t2.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = t2.add_run('미션 의도 파악 스터디 리포트')
r2.font.size = Pt(18)
r2.font.color.rgb = RGBColor(80, 80, 80)

doc.add_paragraph()
p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('2026년 5월 25일 (월) 20:00')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(120, 120, 120)

p = doc.add_paragraph()
p.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p.add_run('정리: 39번 유아영')
r.font.size = Pt(12)
r.font.color.rgb = RGBColor(120, 120, 120)

doc.add_page_break()

# ═══════ 미션 1 ═══════
add_colored_heading('미션 1 — BB 예측', level=1)
doc.add_paragraph()

tag_line('BB 예측', '55번')
doc.add_paragraph()

h = add_colored_heading('분석', level=2)
bullet('', bold_prefix='표면적: ')
p = doc.add_paragraph('정말 좋아했던 경험 기반으로 하는 게 맞을 듯', style='List Bullet 2')

bullet('', bold_prefix='심층적: ')
p = doc.add_paragraph('욕구 기반 접근 / 설득력이 중요', style='List Bullet 2')

doc.add_paragraph()
h = add_colored_heading('추가 인사이트', level=2)
bullet('21기에서 모바일 vs PC 비율도 언급된 적 있음')
bullet('예시: "서울 혼술바" — 지방에서 각잡고 올라오는 것 vs 홍대 근처에서 약속 끝나고 검색하는 것')
bullet('→ 같은 키워드라도 타겟의 상황/맥락이 다르다')

divider()

# ═══════ 미션 2 ═══════
add_colored_heading('미션 2 — 프레임 조종 (타겟 프레임 조종 미션)', level=1)
doc.add_paragraph()

h = add_colored_heading('미션 의도', level=2)
bullet('더블바인드 활용')
bullet('"나한테 있는 프레임" 수준이 아니라 성장과 연결하는 방향으로 써야 함')

doc.add_paragraph()
h = add_colored_heading('BB 예측', level=2)
bullet('', bold_prefix='라온님: ')
doc.add_paragraph('"이혼" 프레임', style='List Bullet 2')
bullet('', bold_prefix='솔지님 (93번): ')
doc.add_paragraph('"여성" 프레임', style='List Bullet 2')
bullet('', bold_prefix='엠군님: ')
doc.add_paragraph('(미공개)', style='List Bullet 2')

divider()

# ═══════ 미션 3 ═══════
add_colored_heading('미션 3 — 욕구 + 사이드 R', level=1)
doc.add_paragraph()

h = add_colored_heading('미션 의도', level=2)
bullet('욕구에만 초점 맞추면 안 됨')
bullet('망가지는 걸 계속 망가지게 두고, 거기서 내 이득만 취하라')
bullet('핵심: "이득의 교집합"을 생각하라')
bullet('"타겟이 갖고 있는 욕구 안에서 이득을 보라"')
bullet('일단 타겟이 하는 걸 무조건 시키고 나서 + 알파')

doc.add_paragraph()
h = add_colored_heading('예시', level=2)
bullet('전자담배 — 100억 번의 영향을 받았을 것')
bullet('고등학생 술 뚫어주는 사례')

doc.add_paragraph()
h = add_colored_heading('BB 예측', level=2)
bullet('', bold_prefix='20번: ')
doc.add_paragraph('음주단속 알림 어플', style='List Bullet 2')
bullet('', bold_prefix='31번(에렌): ')
doc.add_paragraph('스포츠토토 AI 프로그램', style='List Bullet 2')

doc.add_paragraph()
h = add_colored_heading('결과', level=2)
tag_line('PP', '27번, 77번')
tag_line('BB', '16번 (에이즈)')

divider()

# ═══════ 미션 4 ═══════
add_colored_heading('미션 4 — 무의식 조종 미션', level=1)
doc.add_paragraph()

h = add_colored_heading('미션 의도', level=2)
bullet('나에게 마이너스를 끼칠 안 좋은 프레임을 바꿔주기')
bullet('내 이득과 연결되도록 조종하라')
bullet('핵심 질문: "당위성이 있나?"')

doc.add_paragraph()
h = add_colored_heading('예시', level=2)
bullet('엠군님 어머님 물 받으시는 사례')
bullet('→ 엠군님의 이득으로 연결되도록 프레임 전환')

divider()

# ═══════ 미션 5 ═══════
add_colored_heading('미션 5 — 프레임 분리', level=1)
doc.add_paragraph()

h = add_colored_heading('미션 의도', level=2)
bullet('프레임 분리를 "타겟" 기준으로 하라')
bullet('분리만 하면 안 됨 → 분리 후 이득이 되어야 진짜 분리')

doc.add_paragraph()
h = add_colored_heading('21기 재미션 사례', level=2)
bullet('재미션 이유: 분리만 하고 이득 연결을 안 한 사람이 많았음')
para('"여러분들은 프레임 분리를 하라는데 탈출을 했다.\n분리한 다음에 이득이 되어야지 분리가 된 건데,\n무리에서 떨어지라 했더니 무인도로 던져버리면 어떡하냐"', italic=True, color=RGBColor(80, 80, 80))

doc.add_paragraph()
h = add_colored_heading('핵심 포인트', level=2)
bullet('타겟이 "가치를 느끼도록" 하는 게 핵심')
bullet('가치 더하기는 기본')

doc.add_paragraph()
h = add_colored_heading('21기 BB 예시', level=2)
para('"괄사가 아니다. 얼굴 조각칼이다"', bold=True, size=13)
bullet('→ 같은 제품, 프레임만 분리했는데 가치가 완전히 달라짐')

divider()

# ═══════ 미션 6 ═══════
add_colored_heading('미션 6 — 적의 적은 내 편', level=1)
doc.add_paragraph()

h = add_colored_heading('미션 의도', level=2)
bullet('"적의 적은 내 편, 적은 곧 목표다"')
bullet('나와 팬덤의 적을 명확히 하라')
bullet('탄핵봉처럼 내 팬덤의 색도 뚜렷하게 만들어 놓고 → 적을 적어야 함')

doc.add_paragraph()
h = add_colored_heading('예시', level=2)
bullet('파란색 지지자 (정치 팬덤)')
bullet('아나키')

divider()
doc.add_paragraph()

# ═══════ 요약 테이블 ═══════
add_colored_heading('미션별 핵심 키워드 요약', level=1)
doc.add_paragraph()

t = doc.add_table(rows=7, cols=3)
t.style = 'Light Grid Accent 1'

headers = ['미션', '핵심 의도', '키워드']
for i, h in enumerate(headers):
    t.rows[0].cells[i].text = h

data = [
    ['1번', 'BB 예측 + 타겟 맥락 분석', '욕구, 설득력, 모바일/PC 비율'],
    ['2번', '프레임 조종 → 성장 연결', '더블바인드, 성장'],
    ['3번', '욕구 + 사이드R → 이득의 교집합', '욕구, 교집합, +알파'],
    ['4번', '무의식 프레임 → 내 이득 연결', '무의식, 당위성'],
    ['5번', '프레임 분리 → 타겟 가치', '탈출X 분리O, 가치더하기'],
    ['6번', '적의 적은 내 편 → 팬덤 색깔', '적=목표, 팬덤'],
]
for ri, row in enumerate(data):
    for ci, val in enumerate(row):
        t.rows[ri+1].cells[ci].text = val

out = '/Users/kwoneren/.openclaw/workspace/엠타트업22기_4조_미션의도파악_0525.docx'
doc.save(out)
print(f'OK: {out}')
