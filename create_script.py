from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

# 스타일 설정
style = doc.styles['Normal']
font = style.font
font.name = '맑은 고딕'
font.size = Pt(12)

# 제목
title = doc.add_heading('', level=0)
run = title.add_run('타겟 프레임 조종 미션 — 발표 스크립트')
run.font.size = Pt(22)
title.alignment = WD_ALIGN_PARAGRAPH.CENTER

# 부제
sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = sub.add_run('31번 에렌  |  엠타트업 22기')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(100, 100, 100)

doc.add_paragraph()

# 섹션 헬퍼
def add_section(title_text, body_lines):
    h = doc.add_heading(title_text, level=2)
    for line in body_lines:
        if line.startswith('['):
            # 시나리오/연출 지시
            p = doc.add_paragraph()
            run = p.add_run(line)
            run.italic = True
            run.font.color.rgb = RGBColor(120, 120, 120)
            run.font.size = Pt(11)
        else:
            p = doc.add_paragraph(line)
            p.style.font.size = Pt(12)

# ── PART 1: 인트로 ──
add_section('PART 1. 인트로 — 내 프레임을 깨다', [
    '[무대 중앙으로 걸어 나간다. 차분하게.]',
    '',
    '안녕하세요, 31번입니다.',
    '',
    '저는 오늘 발표로 우수 피드백을 받기 위해 나왔습니다.',
    '',
    '그리고 제 안에는 이런 프레임이 하나 있습니다.',
    '',
    '"발표로 우수 피드백을 받으려면 무언가 퍼포먼스를 보여줘야 하고,\n기세 있고 당당하게 가야 한다."',
    '',
    '근데 여러분도 보셔서 아시겠지만...\n저는 기세 있고 재밌고 당당하게 하는 스타일이 아닙니다.',
    '',
    '그래서 오늘, 저는\n"재밌고 기세 있게 해야만 우수 피드백을 받을 수 있다"는\n제 프레임을 깨러 나왔습니다.',
])

# ── PART 2: 청중 프레임 깨기 ──
add_section('PART 2. 청중의 프레임 깨기 — "발표"라는 감옥', [
    '[청중을 천천히 둘러보며]',
    '',
    '여기서 발표 해보신 분, 손 한번만 들어주실래요?',
    '',
    '[손 드는 사람들 확인]',
    '',
    '와... 여기서 발표하려니까 너무 떨리지 않았나요, 여러분?',
    '',
    '네, 근데 그게 — "발표"라는 프레임 안에 갇혀서 그렇습니다.',
    '"발표"라고 생각하면 너무 떨립니다.',
    '',
    '그래서 저는 지금부터 발표가 아니라,\n"31번의 스케치북"이라는 컨셉으로 진행해보겠습니다.',
    '',
    '[한 박자 쉬고, 톤을 편하게 바꾼다]',
    '',
    '우리는 여기 왜 온 거죠?',
    '',
    '...100억 벌기 위해서요.',
    '',
    '맞습니다. 100억이라는 적을 향해서\n같은 편끼리 이야기 나누는 겁니다.',
    '',
    '그렇게 생각하면, 발표가 긴장될 수가 없습니다.\n목표가 같은 사람들끼리 대화하는 건데요.',
])

# ── PART 3: 엠군님 프레임 깨기 ──
add_section('PART 3. 엠군님의 프레임 깨기 — 강의는 에너지를 뺏는다?', [
    '[자연스럽게 엠군님 쪽을 보며]',
    '',
    '그리고 하나 더 깨고 싶은 프레임이 있습니다.',
    '',
    '엠군님한테는 이런 프레임이 있을 수 있어요.',
    '"강의는 힘들고, 에너지가 빠진다."',
    '',
    '근데 만약에...\n강의가 끝나고 좀 쉬신 다음에,\n가벼운 느낌으로 초콜릿 코너를 한번 해주실 수 있다면 어떨까요?',
    '',
    '에너지를 뺏기는 게 아니라,\n오히려 에너지를 얻고 가는 시간.',
    '',
    '그게 가능하다면,\n"강의 = 에너지 소모"라는 프레임도 깨지는 겁니다.',
])

# ── PART 4: 정리 & 클로징 ──
add_section('PART 4. 정리 — 세 개의 프레임', [
    '[차분하게, 정리하듯]',
    '',
    '저는 오늘 이 자리에서\n저와 여러분과 엠군님, 세 사람의 프레임을 깼습니다.',
    '',
    '눈치채셨을까요?',
    '',
    '[손가락으로 하나씩 짚으며]',
    '',
    '하나. 엠군님의 프레임.\n"강의는 에너지가 빠진다" → 오히려 에너지를 얻고 재밌게 할 수도 있다.',
    '',
    '둘. 여러분의 프레임.\n"발표니까 긴장된다" → 목표가 같은 사람끼리 이야기 나누는 건데, 긴장될 게 뭐가 있나.',
    '',
    '셋. 나의 프레임.\n"재밌고 기세 좋아야 우수 피드백을 받는다"',
    '',
    '[살짝 웃으며]',
    '',
    '...이 마지막 프레임은,\n여러분의 소중한 한 표로 깨주시면 감사하겠습니다.',
])

add_section('PART 5. 마무리', [
    '[목소리 힘주며]',
    '',
    '마지막으로, 다 같이 외치고 마무리하겠습니다.',
    '',
    '22기 지백기수!',
    '',
    '[청중과 함께]',
    '',
    '화이팅!!!',
    '',
    '[인사하고 퇴장]',
])

# 하단 메모
doc.add_paragraph()
p = doc.add_paragraph()
run = p.add_run('━━━━━━━━━━━━━━━━━━━━━━━━')
run.font.color.rgb = RGBColor(200, 200, 200)

p = doc.add_paragraph()
run = p.add_run('연출 노트')
run.bold = True
run.font.size = Pt(11)

notes = [
    '전체 톤: 차분하고 진지하게. 퍼포먼스 없이 메시지로 승부.',
    '스케치북 컨셉: 필요시 실제 스케치북/화이트보드에 키워드 적으며 진행 가능.',
    '엠군님 초콜릿 코너 제안: 사전에 양해 구하거나, 즉석 제안으로 서프라이즈 효과.',
    '마지막 "한 표" 부분에서 웃음 유도 — 유일한 유머 포인트, 힘 빼고 자연스럽게.',
    '예상 소요시간: 3~5분',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

out = '/Users/kwoneren/.openclaw/workspace/엠타트업_프레임조종_발표스크립트.docx'
doc.save(out)
print(f'OK: {out}')
